"""Generate Word doc listing SOHModel changes outside carletonwalkingv1."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/mnt/c/Users/doria/Documents/model-soh/carletonwalkingv1/SOHModel_Changes_Outside_Box.docx"

doc = Document()

title = doc.add_heading("SOHModel Changes Outside carletonwalkingv1", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "This document describes modifications made to the shared SOHModel library "
    "(not inside the carletonwalkingv1 simulation box) to support scheduler-based "
    "car evacuation at Carleton campus."
)

doc.add_heading("Context", level=1)
doc.add_paragraph(
    "The Carleton project lives in carletonwalkingv1 and references the shared "
    "SOHModel project via a ProjectReference in SOHCarletonWalkingBox.csproj. "
    "Scheduler-based spawning requires CarDriverSchedulerLayer and related car-driving "
    "logic that already lives in SOHModel. Bug fixes and integration hooks therefore "
    "had to be applied in SOHModel rather than only inside the box."
)

doc.add_heading("Files Modified", level=1)

files = [
    (
        "SOHModel/Car/Model/CarLayer.cs",
        "Skip upfront agent spawn when using a scheduler CSV",
        [
            "Added IsScheduledCarDriverFile() to detect schedule-format CSV files "
            "(headers include startTime and spawningIntervalInMinutes).",
            "In InitLayer(), when the CarDriver agent file is a scheduler CSV, "
            "CarLayer no longer calls AgentManager.SpawnAgents() at startup.",
            "Without this change, cars would be spawned twice: once from CarLayer "
            "at init and again from CarDriverSchedulerLayer on the schedule.",
            "Direct-spawn CSV files (e.g. per-car rows with startLat/destLat) are "
            "unchanged and still spawn through CarLayer as before.",
        ],
    ),
    (
        "SOHModel/Multimodal/Model/CarDriverSchedulerLayer.cs",
        "Spawn CarDriver agents from schedule rows",
        [
            "Implements Schedule() for the MARS scheduler: reads start/destination "
            "coordinates from schedule rows (WKT geometry or startLat/startLon and "
            "destLat/destLon columns).",
            "Creates CarDriver instances with driveMode, trafficCode, and osmRoute "
            "from the schedule row.",
            "Uses a no-op register in the CarDriver constructor, then calls "
            "RegisterAgent() after adding the driver to CarLayer.Driver — avoids "
            "double registration (same pattern as other scheduler layers).",
            "Wraps spawn in try/catch so a bad coordinate or unroutable spawn logs "
            "a message instead of stopping the whole simulation.",
            "ReadCoordinate() falls back to lat/lon columns when geometry is not "
            "present in the schedule row.",
        ],
    ),
    (
        "SOHModel/Car/Model/CarDriver.cs",
        "Unique agent IDs and safer CSV export",
        [
            "Added ID = Guid.NewGuid() in the constructor. BusDriver and TrainDriver "
            "already did this; CarDriver did not. Scheduler spawns use new CarDriver() "
            "directly (not AgentManager.SpawnAgents), so every spawned car kept the "
            "default ID 00000000-0000-0000-0000-000000000000. CarLayer.Driver.TryAdd() "
            "then only kept the first car — the root cause of “1 completed trip” with "
            "the scheduler.",
            "Hardened CurrentEdgeId getter for CSV output. The main config writes "
            "CarDriver.csv each tick; when osmid was missing or null on an edge, "
            "the old code threw NullReferenceException around 6:01 sim time. The getter "
            "now null-checks Car, CurrentEdge, and Attributes and returns \"-1\" safely.",
        ],
    ),
]

for path, summary, bullets in files:
    doc.add_heading(path, level=2)
    p = doc.add_paragraph()
    run = p.add_run("Summary: ")
    run.bold = True
    p.add_run(summary)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

doc.add_heading("What Was Not Changed in SOHModel", level=1)
doc.add_paragraph(
    "No changes were required to Mars.Core, Mars.Components, or other SOH boxes. "
    "Campus-specific configuration, schedules, drive graph, analysis scripts, and "
    "Program.cs wiring remain inside carletonwalkingv1 only."
)

doc.add_heading("Impact on Other Projects", level=1)
doc.add_paragraph(
    "Because SOHModel is shared, these edits affect any solution that references "
    "SOHModel and uses CarDriver with CarDriverSchedulerLayer. The changes are "
    "backward-compatible for direct CSV spawn and add defensive behavior for "
    "scheduler spawn and CSV logging."
)

doc.add_heading("Repository Paths", level=1)
table = doc.add_table(rows=4, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Location"
hdr[1].text = "Path (relative to model-soh repo root)"
rows = [
    ("Simulation box (local)", "carletonwalkingv1/"),
    ("Shared library", "SOHModel/"),
    ("Car layer", "SOHModel/Car/Model/CarLayer.cs"),
    ("Car driver agent", "SOHModel/Car/Model/CarDriver.cs"),
]
for i, (a, b) in enumerate(rows, start=0):
    if i == 0:
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    else:
        if i < len(table.rows):
            table.rows[i].cells[0].text = a
            table.rows[i].cells[1].text = b

# fix table - I created 4 rows but only 4 items - good

doc.add_paragraph()
doc.add_paragraph("Document generated for the Carleton campus evacuation work (carletonwalkingv1).")

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

doc.save(OUT)
print(f"Wrote {OUT}")
